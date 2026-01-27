"""
BhashaAI Backend - Document Service

Handles document upload, text extraction, and processing.
"""

import logging
from io import BytesIO
from typing import Optional
from uuid import UUID

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.storage import get_storage_service
from app.models import Document, DocumentStatus, FileType

logger = logging.getLogger(__name__)


class DocumentService:
    """
    Document service for upload and text extraction.
    """
    
    def __init__(self, db: AsyncSession):
        """Initialize with database session."""
        self.db = db
        self._storage = None
    
    @property
    def storage(self):
        """Lazy-load storage service."""
        if self._storage is None:
            self._storage = get_storage_service()
        return self._storage
    
    async def upload_document(
        self,
        user_id: UUID,
        file_data: bytes,
        filename: str,
        content_type: str,
        institution_id: Optional[UUID] = None,
        subject: Optional[str] = None,
        grade_level: Optional[str] = None,
    ) -> Document:
        """
        Upload a document and create database record.
        
        Args:
            user_id: Owner's user ID
            file_data: File content bytes
            filename: Original filename
            content_type: MIME type
            institution_id: Associated institution
            subject: Subject category
            grade_level: Grade level
        
        Returns:
            Document: Created document record
        """
        # Determine file type
        file_type = self._get_file_type(filename, content_type)
        
        # Upload to MinIO
        file_stream = BytesIO(file_data)
        object_name = self.storage.upload_file(
            file_data=file_stream,
            filename=filename,
            content_type=content_type,
            folder="documents",
        )
        
        # Create document record
        document = Document(
            user_id=str(user_id),
            institution_id=str(institution_id) if institution_id else None,
            filename=filename,
            file_url=object_name,
            file_type=file_type,
            file_size=len(file_data),
            mime_type=content_type,
            subject=subject,
            grade_level=grade_level,
            processing_status=DocumentStatus.PENDING,
        )
        
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)
        
        logger.info(f"Document uploaded: {document.id}")
        return document
    
    async def get_document(
        self,
        document_id: UUID,
        user_id: UUID,
    ) -> Optional[Document]:
        """
        Get document by ID for a specific user.
        
        Args:
            document_id: Document ID
            user_id: User ID for ownership check
        
        Returns:
            Document or None
        """
        stmt = select(Document).where(
            Document.id == str(document_id),
            Document.user_id == str(user_id),
            Document.is_active == True,
        )
        result = await self.db.execute(stmt)
        return result.scalar_one_or_none()
    
    async def list_documents(
        self,
        user_id: UUID,
        page: int = 1,
        per_page: int = 20,
        search: Optional[str] = None,
        file_type: Optional[str] = None,
    ) -> tuple[list[Document], int]:
        """
        List documents for a user with pagination and filtering.
        
        Args:
            user_id: User ID
            page: Page number
            per_page: Items per page
            search: Search term for filename
            file_type: Filter by file type
        
        Returns:
            tuple: (documents, total_count)
        """
        # Base query
        query = select(Document).where(
            Document.user_id == str(user_id),
            Document.is_active == True,
        )

        # Apply filters
        if search:
            query = query.where(Document.filename.ilike(f"%{search}%"))
        
        if file_type:
            # Map simplified types to enum if needed, or rely on frontend passing correct enum values
            # Assuming frontend passes 'pdf', 'docx', 'txt' matches DB enum values or string storage
            # The model uses FileType enum, so we might need to cast if strict
            query = query.where(Document.file_type == file_type)

        # Count total
        # We need to replicate the query for counting to respect filters
        # Or usually separate count query is cleaner
        
        # Optimize: reuse conditions?
        # For simple implementation, rebuild count query or use func.count() on the same query object structure
        
        # Let's simple rebuild for clarity
        count_stmt = select(Document).where(
            Document.user_id == str(user_id),
            Document.is_active == True,
        )
        if search:
            count_stmt = count_stmt.where(Document.filename.ilike(f"%{search}%"))
        if file_type:
            count_stmt = count_stmt.where(Document.file_type == file_type)
            
        count_result = await self.db.execute(count_stmt)
        total = len(count_result.scalars().all())
        
        # Get paginated results
        offset = (page - 1) * per_page
        stmt = (
            query
            .order_by(Document.created_at.desc())
            .offset(offset)
            .limit(per_page)
        )
        result = await self.db.execute(stmt)
        documents = result.scalars().all()
        
        return list(documents), total
    
    async def delete_document(
        self,
        document_id: UUID,
        user_id: UUID,
    ) -> bool:
        """
        Soft delete a document.
        
        Args:
            document_id: Document ID
            user_id: User ID for ownership check
        
        Returns:
            bool: True if deleted
        """
        document = await self.get_document(document_id, user_id)
        if not document:
            return False
        
        document.is_active = False
        await self.db.commit()
        
        logger.info(f"Document deleted: {document_id}")
        return True
    
    async def extract_text(self, document: Document) -> str:
        """
        Extract text content from document.
        
        Args:
            document: Document record
        
        Returns:
            str: Extracted text
        """
        # Download file
        file_data = self.storage.download_file(document.file_url)
        
        # Extract based on file type
        if document.file_type == FileType.PDF:
            return await self._extract_pdf_text(file_data)
        elif document.file_type == FileType.DOCX:
            return await self._extract_docx_text(file_data)
        elif document.file_type == FileType.TXT:
            return file_data.decode("utf-8", errors="ignore")
        else:
            return ""
    
    async def _extract_pdf_text(self, file_data: bytes) -> str:
        """Extract text from PDF."""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(BytesIO(file_data))
            text_parts = []
            
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""
    
    async def _extract_docx_text(self, file_data: bytes) -> str:
        """Extract text from DOCX."""
        try:
            import docx
            
            doc = docx.Document(BytesIO(file_data))
            text_parts = []
            
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    def _get_file_type(self, filename: str, content_type: str) -> FileType:
        """Determine file type from filename and content type."""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        
        if ext == "pdf" or "pdf" in content_type:
            return FileType.PDF
        elif ext in ("docx", "doc") or "document" in content_type:
            return FileType.DOCX
        elif ext == "txt" or "text" in content_type:
            return FileType.TXT
        else:
            return FileType.TXT
    
    def get_download_url(self, document: Document, expires_hours: int = 1) -> str:
        """Get presigned download URL for document."""
        return self.storage.get_presigned_url(document.file_url, expires_hours)

    def download_document_file(self, document: Document) -> bytes:
        """
        Download file content from storage.
        
        Args:
            document: Document object
            
        Returns:
            bytes: File content
        """
        return self.storage.download_file(document.file_url)
